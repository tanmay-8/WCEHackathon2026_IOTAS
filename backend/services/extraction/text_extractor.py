"""
Text Extractor - Extract text content from various file formats.

Supports:
- PDF files (.pdf) - both text-based and scanned
- Word documents (.docx)
- Image files with OCR (.png, .jpg, .jpeg)
- Scanned documents using PaddleOCR (free & open source)

All extracted text flows through existing LLM extraction and ingestion pipeline.
"""

from typing import Tuple, Optional, Dict, Any
import os
import io
from pathlib import Path

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    from paddleocr import PaddleOCR
except ImportError:
    PaddleOCR = None

try:
    from pdf2image import convert_from_path, convert_from_bytes
except ImportError:
    convert_from_path = None
    convert_from_bytes = None


class TextExtractor:
    """
    Multi-format text extraction service.
    Supports PDF, DOCX, and future OCR support for images.
    """
    
    # Supported file types
    SUPPORTED_FORMATS = {
        '.pdf': 'PDF Document',
        '.docx': 'Word Document',
        '.doc': 'Word Document (Legacy)',
        '.txt': 'Plain Text',
        '.png': 'Image (PNG)',
        '.jpg': 'Image (JPEG)',
        '.jpeg': 'Image (JPEG)',
    }
    
    # Chunk size for PDFs (chars) - keeps context windows reasonable
    MAX_CHUNK_SIZE = 1000
    OVERLAP = 100
    
    def __init__(self):
        """Initialize text extractor with OCR support."""
        self.pdf_supported = PdfReader is not None
        self.docx_supported = Document is not None
        self.image_supported = Image is not None
        self.ocr_supported = PaddleOCR is not None
        self.pdf2image_supported = convert_from_bytes is not None
        
        # Initialize PaddleOCR on first use (lazy loading)
        self.ocr = None
        if self.ocr_supported:
            try:
                print("Initializing PaddleOCR... (this may take a moment on first run)")
                self.ocr = PaddleOCR(use_angle_cls=True, lang='en')
                print("PaddleOCR initialized successfully")
            except Exception as e:
                print(f"Warning: Failed to initialize PaddleOCR: {str(e)}")
    
    def extract_from_file(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from a file.
        
        Args:
            file_path: Path to the file to extract from
            
        Returns:
            Tuple of (extracted_text, metadata)
            
        Raises:
            ValueError: If file type is not supported or extraction fails
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise ValueError(f"File not found: {file_path}")
        
        file_ext = file_path.suffix.lower()
        
        if file_ext not in self.SUPPORTED_FORMATS:
            raise ValueError(
                f"Unsupported file format: {file_ext}. "
                f"Supported formats: {', '.join(self.SUPPORTED_FORMATS.keys())}"
            )
        
        file_size = file_path.stat().st_size
        print(f"Extracting text from: {file_path.name} ({file_size} bytes)")
        
        # Route to appropriate extractor
        if file_ext == '.pdf':
            return self._extract_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self._extract_docx(file_path)
        elif file_ext in ['.png', '.jpg', '.jpeg']:
            return self._extract_image_ocr(file_path)
        elif file_ext == '.txt':
            return self._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported format: {file_ext}")
    
    def extract_from_bytes(
        self, 
        file_bytes: bytes, 
        filename: str
    ) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from file bytes (useful for uploaded files).
        
        Args:
            file_bytes: Raw file bytes
            filename: Original filename (for type detection)
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        file_ext = Path(filename).suffix.lower()
        
        if file_ext == '.pdf':
            return self._extract_pdf_bytes(file_bytes, filename)
        elif file_ext in ['.docx', '.doc']:
            return self._extract_docx_bytes(file_bytes, filename)
        elif file_ext in ['.png', '.jpg', '.jpeg']:
            return self._extract_image_ocr_bytes(file_bytes, filename)
        elif file_ext == '.txt':
            return self._extract_txt_bytes(file_bytes, filename)
        else:
            raise ValueError(f"Unsupported format: {file_ext}")
    
    # ── PDF Extraction ──────────────────────────────────────────────────────
    
    def _extract_pdf(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF file."""
        if not self.pdf_supported:
            raise ValueError("PDF support requires PyPDF2. Install: pip install PyPDF2")
        
        try:
            text_pages = []
            metadata = {}
            
            with open(file_path, 'rb') as f:
                pdf_reader = PdfReader(f)
                metadata['total_pages'] = len(pdf_reader.pages)
                
                # Extract from each page
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text()
                    if text.strip():
                        text_pages.append(text)
                
                # Try to extract PDF metadata
                if pdf_reader.metadata:
                    metadata['pdf_title'] = pdf_reader.metadata.get('/Title', '')
                    metadata['pdf_author'] = pdf_reader.metadata.get('/Author', '')
            
            full_text = '\n\n'.join(text_pages)
            metadata['extraction_method'] = 'PyPDF2'
            metadata['format'] = 'PDF'
            metadata['filename'] = file_path.name
            metadata['text_preview'] = full_text[:200]
            
            return full_text, metadata
            
        except Exception as e:
            raise ValueError(f"Failed to extract PDF: {str(e)}")
    
    def _extract_pdf_bytes(self, file_bytes: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF bytes. Falls back to OCR if PDF is scanned."""
        if not self.pdf_supported:
            raise ValueError("PDF support requires PyPDF2. Install: pip install PyPDF2")
        
        try:
            text_pages = []
            metadata = {}
            
            pdf_reader = PdfReader(io.BytesIO(file_bytes))
            metadata['total_pages'] = len(pdf_reader.pages)
            
            # Try text extraction first
            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    text_pages.append(text)
            
            if pdf_reader.metadata:
                metadata['pdf_title'] = pdf_reader.metadata.get('/Title', '')
                metadata['pdf_author'] = pdf_reader.metadata.get('/Author', '')
            
            full_text = '\n\n'.join(text_pages)
            
            # Check if PDF is scanned (very little text extracted)
            if len(full_text.strip()) < 100:  # Less than 100 chars = likely scanned
                print(f"Detected scanned PDF: {filename}. Using OCR...")
                return self._extract_pdf_scanned_bytes(file_bytes, filename, metadata)
            
            metadata['extraction_method'] = 'PyPDF2'
            metadata['format'] = 'PDF'
            metadata['filename'] = filename
            metadata['text_preview'] = full_text[:200]
            metadata['is_scanned'] = False
            
            return full_text, metadata
            
        except Exception as e:
            # If text extraction fails, try OCR
            print(f"Text extraction failed, falling back to OCR: {str(e)}")
            return self._extract_pdf_scanned_bytes(file_bytes, filename, {})
    
    # ── DOCX Extraction ─────────────────────────────────────────────────────
    
    def _extract_docx(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX file."""
        if not self.docx_supported:
            raise ValueError("DOCX support requires python-docx. Install: pip install python-docx")
        
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            full_text = '\n\n'.join(text_parts)
            
            metadata = {
                'extraction_method': 'python-docx',
                'format': 'DOCX',
                'filename': file_path.name,
                'text_preview': full_text[:200],
                'paragraph_count': len([p for p in doc.paragraphs if p.text.strip()]),
                'table_count': len(doc.tables)
            }
            
            return full_text, metadata
            
        except Exception as e:
            raise ValueError(f"Failed to extract DOCX: {str(e)}")
    
    def _extract_docx_bytes(self, file_bytes: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX bytes."""
        if not self.docx_supported:
            raise ValueError("DOCX support requires python-docx. Install: pip install python-docx")
        
        try:
            doc = Document(io.BytesIO(file_bytes))
            
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            full_text = '\n\n'.join(text_parts)
            
            metadata = {
                'extraction_method': 'python-docx',
                'format': 'DOCX',
                'filename': filename,
                'text_preview': full_text[:200],
                'paragraph_count': len([p for p in doc.paragraphs if p.text.strip()]),
                'table_count': len(doc.tables)
            }
            
            return full_text, metadata
            
        except Exception as e:
            raise ValueError(f"Failed to extract DOCX: {str(e)}")
    
    # ── Image OCR Extraction ────────────────────────────────────────────────
    
    def _extract_image_ocr(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from image using OCR (PaddleOCR)."""
        if not self.ocr:
            raise ValueError(
                "OCR support requires PaddleOCR. Install: pip install paddleocr"
            )
        
        try:
            image = Image.open(file_path)
            text = self._run_ocr_on_image(image)
            
            metadata = {
                'extraction_method': 'PaddleOCR',
                'format': 'Image',
                'filename': file_path.name,
                'text_preview': text[:200],
                'image_size': image.size
            }
            
            return text, metadata
            
        except Exception as e:
            raise ValueError(f"Failed to extract text from image: {str(e)}")
    
    def _extract_image_ocr_bytes(self, file_bytes: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from image bytes using OCR (PaddleOCR)."""
        if not self.ocr:
            raise ValueError(
                "OCR support requires PaddleOCR. Install: pip install paddleocr"
            )
        
        try:
            image = Image.open(io.BytesIO(file_bytes))
            text = self._run_ocr_on_image(image)
            
            metadata = {
                'extraction_method': 'PaddleOCR',
                'format': 'Image',
                'filename': filename,
                'text_preview': text[:200],
                'image_size': image.size
            }
            
            return text, metadata
            
        except Exception as e:
            raise ValueError(f"Failed to extract text from image: {str(e)}")
    
    # ── Plain Text Extraction ───────────────────────────────────────────────
    
    def _extract_txt(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            metadata = {
                'extraction_method': 'plain_text',
                'format': 'TXT',
                'filename': file_path.name,
                'text_preview': text[:200],
                'line_count': len(text.split('\n'))
            }
            
            return text, metadata
            
        except Exception as e:
            raise ValueError(f"Failed to extract TXT: {str(e)}")
    
    def _extract_txt_bytes(self, file_bytes: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from plain text bytes."""
        try:
            text = file_bytes.decode('utf-8')
            
            metadata = {
                'extraction_method': 'plain_text',
                'format': 'TXT',
                'filename': filename,
                'text_preview': text[:200],
                'line_count': len(text.split('\n'))
            }
            
            return text, metadata
            
        except Exception as e:
            raise ValueError(f"Failed to extract TXT: {str(e)}")
    
    # ── Scanned PDF OCR Extraction ──────────────────────────────────────────
    
    def _extract_pdf_scanned_bytes(
        self, 
        file_bytes: bytes, 
        filename: str,
        metadata: Dict[str, Any]
    ) -> Tuple[str, Dict[str, Any]]:
        """Extract text from scanned PDF using OCR (PaddleOCR)."""
        if not self.ocr:
            raise ValueError(
                "OCR support requires PaddleOCR. Install: pip install paddleocr pdf2image"
            )
        
        if not convert_from_bytes:
            raise ValueError(
                "PDF to image conversion requires pdf2image. Install: pip install pdf2image"
            )
        
        try:
            # Convert PDF to images
            print(f"Converting PDF to images for OCR: {filename}")
            images = convert_from_bytes(file_bytes)
            
            text_pages = []
            for page_num, image in enumerate(images, 1):
                try:
                    page_text = self._run_ocr_on_image(image)
                    if page_text.strip():
                        text_pages.append(page_text)
                        print(f"  Page {page_num}/{len(images)}: Extracted {len(page_text)} chars")
                except Exception as e:
                    print(f"  Page {page_num}: OCR failed - {str(e)}")
                    continue
            
            full_text = '\n\n'.join(text_pages)
            
            metadata['extraction_method'] = 'PaddleOCR (Scanned PDF)'
            metadata['format'] = 'PDF (Scanned)'
            metadata['filename'] = filename
            metadata['is_scanned'] = True
            metadata['ocr_confidence'] = 'High'
            metadata['text_preview'] = full_text[:200]
            
            if 'total_pages' not in metadata:
                metadata['total_pages'] = len(images)
            
            return full_text, metadata
            
        except Exception as e:
            raise ValueError(f"Failed to extract text from scanned PDF: {str(e)}")
    
    # ── OCR Helper Method ───────────────────────────────────────────────────
    
    def _run_ocr_on_image(self, image: Image.Image) -> str:
        """
        Run PaddleOCR on a PIL image and extract text.
        
        Args:
            image: PIL Image object
            
        Returns:
            Extracted text
        """
        if not self.ocr:
            raise ValueError("PaddleOCR not initialized")
        
        try:
            # Convert PIL image to numpy array
            import numpy as np
            image_array = np.array(image)
            
            # Run OCR
            ocr_result = self.ocr.ocr(image_array, cls=True)
            
            # Extract text from OCR result
            # PaddleOCR returns a list of results per line
            text_lines = []
            for line_result in ocr_result:
                if line_result:
                    for word_info in line_result:
                        text = word_info[1][0]  # Get the text
                        confidence = word_info[1][1]  # Get confidence score
                        
                        # Only include high-confidence results
                        if confidence > 0.3:
                            text_lines.append(text)
            
            return ' '.join(text_lines)
            
        except Exception as e:
            raise ValueError(f"OCR processing failed: {str(e)}")
    
    # ── Chunking Utility ────────────────────────────────────────────────────
    
    @staticmethod
    def chunk_text(
        text: str,
        chunk_size: int = MAX_CHUNK_SIZE,
        overlap: int = OVERLAP
    ) -> list[str]:
        """
        Split text into overlapping chunks.
        
        Useful for processing large documents in the LLM context window.
        
        Args:
            text: Text to chunk
            chunk_size: Size of each chunk (characters)
            overlap: Overlap between chunks
            
        Returns:
            List of text chunks
        """
        chunks = []
        start = 0
        
        while start < len(text):
            end = min(start + chunk_size, len(text))
            chunk = text[start:end]
            chunks.append(chunk)
            
            # Move start position, accounting for overlap
            start = end - overlap
            
            # Avoid infinite loop on very small overlap
            if start <= 0:
                break
        
        return chunks


# Global instance for convenience
text_extractor = TextExtractor()
